<<<<<<< HEAD
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# 1. Function to generate the Detailed Word Document (Based on Source 3)
def generate_cover_letter(data):
    doc = Document()
    
    # Global Font Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Header Section
    doc.add_paragraph("To\nThe Visa Officer\nEmbassy of the Kingdom of Belgium\nIslamabad, Pakistan")
    
    # Subject
    subject = doc.add_paragraph()
    subject.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subject_run = subject.add_run(f"Subject: {data['subject']}")
    subject_run.bold = True
    subject_run.underline = True

    # Introduction
    doc.add_paragraph(
        f"I, {data['full_name']}, son of {data['father_name']}, a Pakistani national, "
        f"holding passport number {data['passport']} and CNIC number {data['cnic']}, "
        f"am applying for a one-year multiple-entry Tourist Visa for travel from "
        f"{data['start_date']} to {data['end_date']}. The purpose of my visit is tourism during this period."
    )

    # Purpose of Travel
    doc.add_heading('Purpose of Travel', level=2)
    doc.add_paragraph(
        f"The purpose of my visit to Belgium is tourism, during which I plan to explore the cultural, "
        f"historical, and architectural highlights of the country through a well-structured travel itinerary. "
        f"My travel is scheduled from {data['start_date']} to {data['end_date']}, during which I will be "
        f"primarily based in Brussels and will also undertake day trips to other major cities."
    )
    
    doc.add_paragraph("In Brussels, I plan to visit:")
    brussels_spots = [
        "Grand Place – iconic central square",
        "Galeries Royales Saint-Hubert – historic luxury arcade",
        "Atomium – unique architectural monument",
        "Mini-Europe, Mont des Arts, and Magritte Museum",
        "Parc du Cinquantenaire – large public park with triumphal arch"
    ]
    for spot in brussels_spots:
        doc.add_paragraph(spot, style='List Bullet')

    doc.add_paragraph("My itinerary also includes day trips to:")
    day_trips = ["Bruges – 'Venice of the North'", "Ghent – Gravensteen Castle", "Antwerp – Diamond district", "Dinant – Citadel of Dinant"]
    for trip in day_trips:
        doc.add_paragraph(trip, style='List Bullet')

    doc.add_paragraph(f"After Belgium, I will travel to the USA from 22 June to 27 June 2026, as I hold a valid US visa and a 10-year UK visa.")

    # Accommodation Details
    doc.add_heading('Accommodation Details', level=2)
    doc.add_paragraph(f"Hotel: {data['hotel_name']}\nAddress: {data['hotel_address']}\nContact: {data['hotel_contact']}")

    # Financial Arrangements
    doc.add_heading('Financial Arrangements', level=2)
    doc.add_paragraph("All travel expenses will be self-funded. Attached are my bank statements (reflecting financial stability) and tax returns for 2024-2025.")

    # Business & Professional Background
    doc.add_heading('Business & Professional Background', level=2)
    doc.add_paragraph(
        f"I am the owner of {data['business_name']}, a construction enterprise registered with FBR "
        f"(NTN: {data['ntn']}) since {data['reg_date']}. I am an active member of {data['chamber']} "
        f"(Membership No: {data['membership']})."
    )

    # Family & Ties
    doc.add_heading('Family & Ties to Pakistan', level=2)
    doc.add_paragraph(
        f"I am a married individual and a father of {data['children']} children, residing in {data['city']}. "
        f"My business and family commitments ensure my return to Pakistan."
    )

    # Travel History
    doc.add_heading('Travel History', level=2)
    doc.add_paragraph(
        "I have extensive travel history including Saudi Arabia (3 visits), Malaysia (2), Turkey (2), "
        "UK (2), Egypt (2), and single visits to USA, Belgium, Spain, Greece, Japan, and others. "
        "I have always complied with visa regulations."
    )

    # Declaration & Request
    doc.add_heading('Declaration & Request', level=2)
    doc.add_paragraph("I affirm my commitment to strictly adhering to all Belgian laws and Schengen regulations. I kindly request the issuance of a short-stay tourist visa.")

    # Closing
    doc.add_paragraph("\nYour Sincerely,")
    doc.add_paragraph(f"{data['full_name']}\nContact: {data['contact']}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 2. Streamlit UI
st.set_page_config(page_title="Visa Pro Generator", layout="wide")
st.title("📄 Professional Visa Cover Letter Generator")
st.info("Source 3 Format: Detailed Itinerary & Travel History included.")

with st.form("detailed_form"):
    c1, c2 = st.columns(2)
    with c1:
        full_name = st.text_input("Full Name", "Rana Muhammad Asif")
        father_name = st.text_input("Father's Name", "Rana Abdul Majeed")
        passport = st.text_input("Passport No", "BJ1880214")
        cnic = st.text_input("CNIC No", "34401-44150217")
    with c2:
        biz_name = st.text_input("Business Name", "RANA BUILDERS")
        ntn = st.text_input("NTN No", "2815574-2")
        chamber = st.text_input("Chamber", "Islamabad Chamber of Commerce & Industry")
        membership = st.text_input("Membership No", "AM-5984")
    
    st.markdown("---")
    submitted = st.form_submit_button("Generate Complete Document")

if submitted:
    # Full data mapping from your source
    final_data = {
        "subject": "Application for Belgium Tourist Visa",
        "full_name": full_name, "father_name": father_name, "passport": passport, "cnic": cnic,
        "start_date": "13 June 2026", "end_date": "22 June 2026",
        "hotel_name": "Dolce by Wyndham La Hulpe Brussels",
        "hotel_address": "135 Chaussée de Bruxelles, 1310 La Hulpe, Belgium",
        "hotel_contact": "+32 2 290 9800",
        "business_name": biz_name, "ntn": ntn, "reg_date": "15 December 2006",
        "chamber": chamber, "membership": membership,
        "children": "two", "city": "Islamabad", "contact": "+92 3007799500"
    }
    
    file_bytes = generate_cover_letter(final_data)
    st.success("Zabardast! Aapka detailed cover letter (Source 3) tayyar hai.")
    st.download_button(
        label="📥 Download Detailed Cover Letter",
        data=file_bytes,
        file_name=f"Visa_Letter_{full_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
=======
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def generate_cover_letter(data):
    doc = Document()
    
    # Global Font Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # --- Header ---
    header = doc.add_paragraph("To\nThe Visa Officer\nEmbassy of the Kingdom of Belgium\nIslamabad, Pakistan")
    header.paragraph_format.space_after = Pt(12)
    
    # --- Subject ---
    subject = doc.add_paragraph()
    subject_run = subject.add_run("Subject: Application for Belgium Tourist Visa")
    subject_run.bold = True
    subject_run.underline = True
    subject.paragraph_format.space_after = Pt(18)

    # --- Introduction ---
    intro = doc.add_paragraph(
        f"I, {data['full_name']}, son of {data['father_name']}, a Pakistani national, "
        f"holding passport number {data['passport']} and CNIC number {data['cnic']}, "
        f"am applying for a tourist visa for travel from {data['start_date']} to {data['end_date']}. "
        f"The purpose of my visit is tourism during this period."
    )
    intro.paragraph_format.space_after = Pt(12)

    # --- Section Styling Helper ---
    def add_section_heading(text):
        h = doc.add_heading(text, level=2)
        run = h.runs[0]
        run.font.color.rgb = RGBColor(79, 129, 189) # Blue color like screenshot
        run.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    # --- Purpose of Travel & Detailed Itinerary ---
    add_section_heading('Purpose of Travel')
    doc.add_paragraph(
        "The purpose of my visit to Belgium is tourism, during which I plan to explore the cultural, "
        "historical, and architectural highlights of the country through a well-structured travel itinerary. "
        f"My travel is scheduled from {data['start_date']} to {data['end_date']}, during which I will be primarily "
        "based in Brussels and will also undertake day trips to other major cities."
    )
    
    doc.add_paragraph("In Brussels, I plan to visit:")
    brussels_points = [
        "Grand Place – iconic central square",
        "Galeries Royales Saint-Hubert – historic luxury arcade",
        "Atomium – unique architectural monument",
        "Mini-Europe, Mont des Arts, and Magritte Museum",
        "Parc du Cinquantenaire – large public park with triumphal arch"
    ]
    for p in brussels_points:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_paragraph("My itinerary also includes day trips to:")
    day_trips = [
        "Bruges – 'Venice of the North'",
        "Ghent – Gravensteen Castle",
        "Antwerp – Diamond district",
        "Dinant – Citadel of Dinant"
    ]
    for p in day_trips:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_paragraph(f"After Belgium, I will travel to {data['next_dest']} as per my travel plans.")

    # --- Accommodation Details ---
    add_section_heading('Accommodation Details')
    doc.add_paragraph(
        f"Hotel: {data['hotel_name']}\n"
        f"Address: {data['hotel_address']}\n"
        f"Contact: {data['hotel_contact']}"
    )

    # --- Financial Arrangements ---
    add_section_heading('Financial Arrangements')
    doc.add_paragraph(
        "All travel expenses will be self-funded. Attached are my bank statements (reflecting financial "
        "stability) and tax returns for the recent years."
    )

    # --- Business & Professional Background ---
    add_section_heading('Business & Professional Background')
    doc.add_paragraph(
        f"I am the owner of {data['business_name']}, a construction enterprise registered with FBR "
        f"(NTN: {data['ntn']}). I am an active member of the Islamabad Chamber of Commerce & Industry."
    )

    # --- Family & Ties to Pakistan ---
    add_section_heading('Family & Ties to Pakistan')
    doc.add_paragraph(data['family_details'])

    # --- Travel History ---
    add_section_heading('Travel History')
    doc.add_paragraph(data['travel_history'])

    # --- Declaration & Request ---
    add_section_heading('Declaration & Request')
    doc.add_paragraph(
        "I affirm my commitment to strictly adhering to all Belgian laws and Schengen regulations. "
        "I kindly request the issuance of a short-stay tourist visa."
    )

    # --- Closing ---
    doc.add_paragraph("\nYour Sincerely,")
    final = doc.add_paragraph(f"{data['full_name']}\nContact: {data['contact']}")
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Streamlit UI ---
st.set_page_config(page_title="Pro Visa Gen", layout="wide")
st.title("📄 Ali Baba Travel Advisor - Professional Cover Letter")

with st.form("visa_form"):
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("👤 Personal Details")
        name = st.text_input("Full Name", "Rana Muhammad Asif")
        fname = st.text_input("Father Name", "Rana Abdul Majeed")
        passp = st.text_input("Passport No", "BJ1880214")
        cnic_no = st.text_input("CNIC No", "34401-44150217")
        biz = st.text_input("Business Name", "RANA BUILDERS")
        ntn_no = st.text_input("NTN", "2815574-2")

    with col2:
        st.subheader("✈️ Travel Details")
        s_date = st.text_input("Start Date", "13 June 2026")
        e_date = st.text_input("End Date", "22 June 2026")
        next_d = st.text_input("Next Destination (if any)", "the USA from 22 June to 27 June 2026")
        h_name = st.text_input("Hotel", "Dolce by Wyndham La Hulpe Brussels")
        h_addr = st.text_area("Hotel Address", "135 Chaussée de Bruxelles, 1310 La Hulpe, Belgium")
        h_cont = st.text_input("Hotel Contact", "+32 2 290 9800")

    st.subheader("📋 Additional Information")
    fam = st.text_area("Family Ties", "I am a married individual and a father of two children, residing in Islamabad. My business and family commitments ensure my return.")
    hist = st.text_area("Detailed Travel History", "Saudi Arabia (3 visits), Malaysia (2), Turkey (2), UK (2), Egypt (2), and single visits to USA, Belgium, Spain, Greece, Japan.")
    
    submit = st.form_submit_button("Generate Long-Form Letter")

if submit:
    user_data = {
        "full_name": name, "father_name": fname, "passport": passp, "cnic": cnic_no,
        "start_date": s_date, "end_date": e_date, "business_name": biz, "ntn": ntn_no,
        "hotel_name": h_name, "hotel_address": h_addr, "hotel_contact": h_cont,
        "family_details": fam, "travel_history": hist, "next_dest": next_d,
        "contact": "+92 3007799500"
    }
    file = generate_cover_letter(user_data)
    st.success("Professional Letter Generated Successfully!")
    st.download_button("📥 Download Final Document", file, f"Visa_Letter_{name}.docx")
>>>>>>> cc2adcf (Initial commit: AI Visa Cover Letter Generator)
